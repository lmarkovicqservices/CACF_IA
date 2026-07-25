"""Convierte un markdown simple (títulos, negritas) a un .docx legible.

Pensado para respuestas_para_revision.md: encabezados '#'/'##' y texto con
**negrita**. No es un parser de markdown completo, alcanza para este caso.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]

BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def agregar_parrafo_con_negritas(doc: Document, texto: str) -> None:
    parrafo = doc.add_paragraph()
    pos = 0
    for match in BOLD_PATTERN.finditer(texto):
        if match.start() > pos:
            parrafo.add_run(texto[pos:match.start()])
        parrafo.add_run(match.group(1)).bold = True
        pos = match.end()
    if pos < len(texto):
        parrafo.add_run(texto[pos:])


def markdown_a_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    for linea in md_path.read_text(encoding="utf-8").splitlines():
        linea = linea.rstrip()
        if not linea:
            continue
        if linea.startswith("## "):
            doc.add_heading(linea[3:], level=2)
        elif linea.startswith("# "):
            doc.add_heading(linea[2:], level=1)
        else:
            agregar_parrafo_con_negritas(doc, linea)

    doc.save(str(docx_path))


def main() -> None:
    if len(sys.argv) != 2:
        print("Uso: python export_docx.py <archivo.md>")
        sys.exit(1)

    md_path = Path(sys.argv[1])
    if not md_path.is_absolute():
        md_path = PROJECT_ROOT / md_path
    if not md_path.exists():
        print(f"No se encontró {md_path}")
        sys.exit(1)

    docx_path = md_path.with_suffix(".docx")
    markdown_a_docx(md_path, docx_path)
    print(f"Generado: {docx_path}")


if __name__ == "__main__":
    main()
